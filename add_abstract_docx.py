import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    # Buka dokumen yang sudah ada
    out_path = r"c:\Users\msi\Documents\New folder\Nickel-Spatiotemporal-GTWR-Leontief\Tesis_Ekstensif_CRISPDM_Nikel_GTWR.docx"
    
    if not os.path.exists(out_path):
        print("Dokumen tidak ditemukan.")
        return

    doc = Document(out_path)

    # Kita ingin menyisipkan Abstrak di awal dokumen, setelah halaman judul.
    # Karena menyisipkan di awal menggunakan python-docx agak rumit tanpa merusak struktur,
    # kita akan menulis ulang dokumen dengan menambahkan Abstrak.
    
    new_doc = Document()
    style = new_doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Fungsi pembantu
    def add_heading_center(text, level):
        h = new_doc.add_heading(text, level)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return h

    def add_padded_paragraph(text, italic=False):
        p = new_doc.add_paragraph()
        if italic:
            p.add_run(text).italic = True
        else:
            p.add_run(text)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # --- COVER PAGE ---
    new_doc.add_paragraph('\n' * 5)
    add_heading_center('NASKAH AKADEMIK EKSTENSIF', 0)
    add_heading_center('EVALUASI SPASIO-TEMPORAL DAMPAK EKONOMI KAWASAN INDUSTRI NIKEL', 1)
    add_heading_center('Integrasi Input-Output Leontief, Pra-pemrosesan Data Resolusi Tinggi, dan Pemodelan GTWR', 2)
    new_doc.add_paragraph('\n' * 3)
    p_author = new_doc.add_paragraph('Disusun oleh: Anwar Rohmadi\nAnalisis Komprehensif Skala Ekstensif\n\n')
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_doc.add_page_break()

    # --- ABSTRAK ---
    add_heading_center('ABSTRAK', 1)
    add_padded_paragraph('Transisi energi global menuju kendaraan listrik memicu ekspansi masif fasilitas pengolahan (smelter) nikel di Indonesia. Evaluasi dampak ekonomi yang ada selama ini bertumpu pada indikator makroekonomi agregat, yang cenderung mengaburkan eksternalitas negatif berupa degradasi Produk Domestik Regional Bruto (PDRB) sektor primer seperti pertanian dan perikanan. Penelitian ini bertujuan untuk mengkuantifikasi klasterisasi kerugian ekonomi tersebut menggunakan pendekatan Data Science dan Ekonometrika Spasial beresolusi tinggi. Metodologi penelitian merangkum distilasi 57 tabel data sekunder menjadi dataset terpusat tingkat fasilitas (Facility-Level). Variabel target kerugian diekstraksi dari matriks Input-Output Leontief, yang kemudian diintegrasikan dengan parameter kapasitas smelter dan Pembangkit Listrik Tenaga Uap (PLTU) captive. Untuk menanggulangi jebakan multikolinearitas dan data leakage pada agregasi provinsi (N=38), observasi direstrukturisasi menjadi Data Panel Spasio-Temporal (N=212). Empat arsitektur algoritma dievaluasi menggunakan skema 5-Fold Cross Validation: Support Vector Regression (SVR), Random Forest, Spatio-Temporal Graph Convolutional Network (ST-GCN), dan Geographically and Temporally Weighted Regression (GTWR). Hasil pengujian empiris membuktikan bahwa GTWR merupakan model paling presisi dengan akurasi Out-of-Sample tertinggi sebesar 88.22%, secara signifikan mengungguli SVR (86.31%) dan Random Forest (86.09%). Keunggulan GTWR menegaskan berlakunya Hukum Geografi Pertama Tobler; bahwa letak spasial dan jarak Euclidean antar smelter menghasilkan efek limpahan (spillover) eksponensial terhadap kerugian ekonomi. Penelitian ini merekomendasikan reorientasi kebijakan investasi melalui moratorium zonasi spasial dan penerapan pajak karbon bergradasi geografis.')
    
    p_keys = new_doc.add_paragraph()
    p_keys.add_run('Kata Kunci: ').bold = True
    p_keys.add_run('Hilirisasi Nikel, Input-Output Leontief, Spatio-Temporal Panel, Geographically and Temporally Weighted Regression, Efek Limpahan.')
    new_doc.add_page_break()

    # Salin elemen dari dokumen asli (mulai dari BAB 1)
    started = False
    for para in doc.paragraphs:
        if "BAB I: PENDAHULUAN" in para.text:
            started = True
        
        if started:
            # Replikasi format Heading jika itu heading
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                new_doc.add_heading(para.text, level)
            else:
                p = new_doc.add_paragraph(para.text)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Salin tabel jika ada
    # (Dokumen 30 halaman memiliki tabel di BAB V)
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = 'Table Grid'
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.cell(i, j).text = cell.text

    # Simpan timpa dokumen lama
    new_doc.save(out_path)
    print("Abstrak berhasil disisipkan.")

if __name__ == '__main__':
    main()
